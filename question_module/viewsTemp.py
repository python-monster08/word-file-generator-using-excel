from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import UploadFileForm
import pandas as pd
from django.db import transaction
import os 
from docx import Document
from docx.shared import Inches
import io
from .models import Question  # Ensure this is the correct path to your models
from datetime import datetime
from django.utils.html import strip_tags
import PIL.Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# def clean_text(text):
#     """
#     Clean text for writing to Word document: replace newlines with spaces and strip HTML tags.
#     """
#     if text:
#         return strip_tags(text.replace('\n', ' ')).strip()
#     return ''

def clean_text(text):
    """
    Clean text for writing to Word document: replace newlines with spaces and strip HTML tags.
    This function ensures that all newlines within the text are replaced with space characters
    to prevent unwanted line breaks in the Word document. HTML tags are removed to maintain 
    the purity of the text.
    """
    text = str(text)
    if text:
        # Replace all newline characters with space
        text = text.replace('\n', ' ')
        # Strip HTML tags
        cleaned_text = strip_tags(text)
        # Strip leading/trailing whitespace
        cleaned_text = cleaned_text.strip()
        return cleaned_text
    return ''



def success_view(request):
    return render(request, 'success.html')


# def upload_file_view(request):
#     if request.method == 'POST':
#         form = UploadFileForm(request.POST, request.FILES)
#         if form.is_valid():
#             file = request.FILES['file']
#             df = pd.read_excel(file)
#             # Replace NaN values with appropriate defaults
#             df.fillna({
#                 'Assertion (A)': '',
#                 'Reason (R)': '',
#                 'List - I Heading': '',
#                 'List - II Heading': '',
#                 'List - I Option A': '',
#                 'List - I Option B': '',
#                 'List - I Option C': '',
#                 'List - I Option D': '',
#                 'List - II Option 1': '',
#                 'List - II Option 2': '',
#                 'List - II Option 3': '',
#                 'List - II Option 4': '',
#                 'List - II Option 5': '',
#                 'Option A': '',
#                 'Option B': '',
#                 'Option C': '',
#                 'Option D': '',
#                 'List - I Selection Text ': '',
#                 'Correct Answer Option': '',
#                 'Correct Answer Discription': '',
#                 'Question Type': '',
#                 'Marks': 0,  # Assuming 0 is an acceptable default for Marks
#                 'Negative Marks': 0  # Assuming 0 is an acceptable default for Negative Marks
#             }, inplace=True)

#             with transaction.atomic():
#                 for index, row in df.iterrows():
#                     Question.objects.create(
#                         question_no=row['Question No'],
#                         question_text=row['Question Text'],
#                         year=row['Year'],
#                         assertion=row['Assertion (A)'],
#                         reason=row['Reason (R)'],
#                         list_i_heading=row['List - I Heading'],
#                         list_ii_heading=row['List - II Heading'],
#                         list_i_option_a=row['List - I Option A'],
#                         list_i_option_b=row['List - I Option B'],
#                         list_i_option_c=row['List - I Option C'],
#                         list_i_option_d=row['List - I Option D'],
#                         list_ii_option_1=row['List - II Option 1'],
#                         list_ii_option_2=row['List - II Option 2'],
#                         list_ii_option_3=row['List - II Option 3'],
#                         list_ii_option_4=row['List - II Option 4'],
#                         list_ii_option_5=row['List - II Option 5'],
#                         option_a=row['Option A'],
#                         option_b=row['Option B'],
#                         option_c=row['Option C'],
#                         option_d=row['Option D'],
#                         list_i_selection_text=row['List - I Selection Text '],
#                         correct_answer_option=row['Correct Answer Option'],
#                         correct_answer_description=row['Correct Answer Discription'],
#                         question_type=row['Question Type'],
#                         marks=row['Marks'],
#                         negative_marks=row['Negative Marks']
#                     )
#             return redirect('success_view')  # Redirect after POST
#     else:
#         form = UploadFileForm()  # An unbound form

#     return render(request, 'upload.html', {'form': form})


def upload_file_view(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES['file']
            df = pd.read_excel(file)

            # Print DataFrame headers for debugging
            print("DataFrame Headers:", df.columns.tolist())

            # Assuming all necessary fields are included in df, otherwise they will be set to an empty string 
            expected_headers = [
                'Question No', 'Question Text', 'Year', 'Assertion (A)', 'Reason (R)',
                'List - I Heading', 'List - II Heading', 'List - I Option A', 'List - I Option B',
                'List - I Option C', 'List - I Option D', 'List - II Option 1', 'List - II Option 2',
                'List - II Option 3', 'List - II Option 4', 'List - II Option 5', 'Option A',
                'Option B', 'Option C', 'Option D', 'Option E', 'List - I Selection Text ',
                'Correct Answer Option', 'Correct Answer Discription', 'Question Type', 'Marks',
                'Negative Marks', 'Table Head A', 'Table Head B', 'Table Head C', 'Table Head D',
                'Head A Data1', 'Head A Data2', 'Head A Data3', 'Head A Data4', 'Head B Data1',
                'Head B Data2', 'Head B Data3', 'Head B Data4', 'Head C Data1', 'Head C Data2',
                'Head C Data3', 'Head C Data4', 'Head D Data1', 'Head D Data2', 'Head D Data3',
                'Head D Data4'
            ]
            default_values = {header: '' for header in expected_headers}
            default_values.update({'Marks': 0, 'Negative Marks': 0})  # Update with defaults for numeric fields
            df.fillna(default_values, inplace=True)

            # Check if the DataFrame contains all the expected headers
            missing_headers = set(expected_headers) - set(df.columns)
            if missing_headers:
                return HttpResponse(f"The following required columns are missing from the Excel file: {missing_headers}", status=400)

            with transaction.atomic():
                for index, row in df.iterrows():
                    Question.objects.create(
                        question_no=row.get('Question No', ''),
                        question_text=clean_text(row.get('Question Text', '')),
                        year=row.get('Year', ''),
                        assertion=row.get('Assertion (A)', ''),
                        reason=row.get('Reason (R)', ''),
                        list_i_heading=clean_text(row.get('List - I Heading', '')),
                        list_ii_heading=clean_text(row.get('List - II Heading', '')),
                        list_i_option_a=clean_text(row.get('List - I Option A', '')),
                        list_i_option_b=clean_text(row.get('List - I Option B', '')),
                        list_i_option_c=clean_text(row.get('List - I Option C', '')),
                        list_i_option_d=clean_text(row.get('List - I Option D', '')),
                        list_ii_option_1=clean_text(row.get('List - II Option 1', '')),
                        list_ii_option_2=clean_text(row.get('List - II Option 2', '')),
                        list_ii_option_3=clean_text(row.get('List - II Option 3', '')),
                        list_ii_option_4=clean_text(row.get('List - II Option 4', '')),
                        list_ii_option_5=clean_text(row.get('List - II Option 5', '')),
                        option_a=clean_text(row.get('Option A', '')),
                        option_b=clean_text(row.get('Option B', '')),
                        option_c=clean_text(row.get('Option C', '')),
                        option_d=clean_text(row.get('Option D', '')),
                        option_e=clean_text(row.get('Option E', '')),
                        list_i_selection_text=clean_text(row.get('List - I Selection Text ', '')),
                        correct_answer_option=clean_text(row.get('Correct Answer Option', '')),
                        correct_answer_description=clean_text(row.get('Correct Answer Discription', '')),
                        question_type=clean_text(row.get('Question Type', '')),
                        marks=row.get('Marks', 0),
                        negative_marks=row.get('Negative Marks', 0),
                        table_head_a=clean_text(row.get('Table Head A', '')),
                        table_head_b=clean_text(row.get('Table Head B', '')),
                        table_head_c=clean_text(row.get('Table Head C', '')),
                        table_head_d=clean_text(row.get('Table Head D', '')),
                        head_a_data1=clean_text(row.get('Head A Data1', '')),
                        head_a_data2=clean_text(row.get('Head A Data2', '')),
                        head_a_data3=clean_text(row.get('Head A Data3', '')),
                        head_a_data4=clean_text(row.get('Head A Data4', '')),
                        head_b_data1=clean_text(row.get('Head B Data1', '')),
                        head_b_data2=clean_text(row.get('Head B Data2', '')),
                        head_b_data3=clean_text(row.get('Head B Data3', '')),
                        head_b_data4=clean_text(row.get('Head B Data4', '')),
                        head_c_data1=clean_text(row.get('Head C Data1', '')),
                        head_c_data2=clean_text(row.get('Head C Data2', '')),
                        head_c_data3=clean_text(row.get('Head C Data3', '')),
                        head_c_data4=clean_text(row.get('Head C Data4', '')),
                        head_d_data1=clean_text(row.get('Head D Data1', '')),
                        head_d_data2=clean_text(row.get('Head D Data2', '')),
                        head_d_data3=clean_text(row.get('Head D Data3', '')),
                        head_d_data4=clean_text(row.get('Head D Data4', ''))
                    )
            return redirect('success_view')  # Redirect after POST
    else:
        form = UploadFileForm()  # An unbound form

    return render(request, 'upload.html', {'form': form})


def set_no_border(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['left', 'top', 'right', 'bottom']:
        border_tag = OxmlElement(f'w:{border}')
        border_tag.set(qn('w:val'), 'nil')  # Set border value to 'nil' to remove it
        border_tag.set(qn('w:sz'), '0')     # Set border size to '0' for good measure
        border_tag.set(qn('w:space'), '0')  # Set border spacing to '0'
        tcBorders.append(border_tag)
    tcPr.append(tcBorders)


# def generate_questions_document(request):
#     try:
#         # Setup directory and document file
#         base_dir = r'C:\Users\HP\Desktop\New Project\hia_project\media\word_files'
#         os.makedirs(base_dir, exist_ok=True)
#         today = datetime.today().strftime('%Y-%m-%d')
#         file_name = f'all_questions_{today}.docx'
#         file_path = os.path.join(base_dir, file_name)
#         document = Document()

#         for question in Question.objects.all():
#             # Create main table for the question and set styles
#             table = document.add_table(rows=0, cols=3)
#             table.style = 'Table Grid'

#             # Handle question text or assertion and reason
#             if question.question_text.strip():
#                 question_text = question.question_text
#             else:
#                 question_text = f"{question.assertion}\n{question.reason}"

#             # Add question info to the main table
#             q_row = table.add_row().cells
#             q_row[0].text = 'Question'

#             if question.list_i_selection_text:
#                 # Handle list_i_selection_text or list-I type questions
#                 q_row[1].text = f"{clean_text(question_text)}\n"
#                 options_format = "\n".join([
#                     f"{i}. {clean_text(getattr(question, f'list_i_option_{chr(96+i)}', ''))}" for i in range(1, 5)
#                     if getattr(question, f'list_i_option_{chr(96+i)}', '')
#                 ])
#                 q_row[1].text += options_format + "\n" + clean_text(question.list_i_selection_text)

#             elif question.list_i_heading and question.list_ii_heading:
#                 # Create a sub-table to handle complex question structures
#                 sub_table = document.add_table(rows=1, cols=2)
#                 sub_table.style = 'Table Grid'

#                 # Apply no border style to all cells in the sub-table as they are created
#                 for cell in sub_table._cells:
#                     set_no_border(cell)

#                 sub_hdr_cells = sub_table.rows[0].cells
#                 sub_hdr_cells[0].text = clean_text(question.list_i_heading)
#                 sub_hdr_cells[1].text = clean_text(question.list_ii_heading)

#                 # Populate sub-table with list options
#                 for i in range(1, 6):
#                     row_cells = sub_table.add_row().cells
#                     for cell in row_cells:
#                         set_no_border(cell)  # Ensure borders are removed for new cells too
#                     li_option = getattr(question, f'list_i_option_{chr(96+i)}', '')
#                     ii_option = getattr(question, f'list_ii_option_{i}', '')
#                     row_cells[0].text = f"{chr(64+i)}. {clean_text(li_option)}" if li_option else ""
#                     row_cells[1].text = f"{i}. {clean_text(ii_option)}" if ii_option else ""

#                 # Clear the original cell content, insert question text and sub-table
#                 q_row[1]._element.clear_content()
#                 p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
#                 # p.add_run(clean_text(question.question_text) + "\n").bold = True
#                 p.add_run(clean_text(question.question_text) + "\n")
#                 q_row[1]._element.append(sub_table._element)

#                 # Add 'Codes:' text below the sub-table within the same cell
#                 p = q_row[1].add_paragraph()
#                 # p.add_run("\nCodes:").bold = True
#                 p.add_run("\nCodes:")


#             else:
#                 # Standard question text handling
#                 q_row[1].text = clean_text(question_text)

#             # Merging cells for question text and image
#             q_row[1].merge(q_row[2])
            
#             # Handling image insertion if available
#             if question.image:
#                 image_path = question.image.path
#                 pil_img = PILImage.open(image_path)
#                 img_io = io.BytesIO()
#                 pil_img.save(img_io, 'JPEG')
#                 img_io.seek(0)
#                 paragraph = q_row[1].add_paragraph()
#                 run = paragraph.add_run()
#                 run.add_picture(img_io, width=Inches(1.5))

#             # Additional rows for type, options, solution, and marks
#             type_row = table.add_row().cells
#             type_row[0].text = 'Type'
#             type_row[1].text = question.question_type
#             type_row[1].merge(type_row[2])

#             correct_option_text = ""  # Store text of the correct answer
#             for opt in ['a', 'b', 'c', 'd']:
#                 option_text = getattr(question, f"option_{opt}", None)
#                 if option_text:
#                     opt_row = table.add_row().cells
#                     opt_row[0].text = 'Option'
#                     opt_row[1].text = f"{opt.upper()}. {clean_text(option_text)}"
#                     opt_row[2].text = 'correct' if opt == question.correct_answer_option.lower() else 'incorrect'
#                     if opt == question.correct_answer_option.lower():
#                         correct_option_text = clean_text(option_text)

#             solution_row = table.add_row().cells
#             solution_row[0].text = 'Solution'
#             solution_row[1].text = correct_option_text
#             solution_row[1].merge(solution_row[2])

#             marks_row = table.add_row().cells
#             marks_row[0].text = 'Marks'
#             marks_row[1].text = str(question.marks)
#             marks_row[2].text = str(question.negative_marks)

#             document.add_paragraph()  # Add space between questions

#         document.save(file_path)
#         return HttpResponse(f"Document created successfully at: {file_path}")
#     except Exception as e:
#         return HttpResponse(f"An error occurred: {str(e)}", status=500)
    


def generate_questions_document(request):
    try:
        # Setup directory and document file to save generated word file
        base_dir = r'C:\Users\HP\Desktop\ClassPlus Word File Creator Project\hia_project\media\word_files'
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'all_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in Question.objects.all():
            # Create main table for the question and set styles
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # Handle question text or assertion and reason
            if question.question_text.strip():
                question_text = question.question_text
            else:
                question_text = f"{clean_text(question.assertion)}\n{clean_text(question.reason)}"
                print(question_text)


            # Add question info to the main table
            q_row = table.add_row().cells
            q_row[0].text = 'Question'

            if question.list_i_selection_text:
                # Handle list_i_selection_text or list-I type questions
                q_row[1].text = f"{question_text}\n"
                options_format = "\n".join([
                    f"{i}. {getattr(question, f'list_i_option_{chr(96+i)}', '')}" for i in range(1, 5)
                    if getattr(question, f'list_i_option_{chr(96+i)}', '')
                ])
                q_row[1].text += options_format # + "\n" + clean_text(question.list_i_selection_text)

            elif question.list_i_heading and question.list_ii_heading:
                # Create a sub-table to handle complex question structures
                sub_table = document.add_table(rows=1, cols=2)
                sub_table.style = 'Table Grid'

                # Apply no border style to all cells in the sub-table as they are created
                for cell in sub_table._cells:
                    set_no_border(cell)

                sub_hdr_cells = sub_table.rows[0].cells
                sub_hdr_cells[0].text = question.list_i_heading
                sub_hdr_cells[1].text = question.list_ii_heading

                # Populate sub-table with list options
                for i in range(1, 6):
                    row_cells = sub_table.add_row().cells
                    for cell in row_cells:
                        set_no_border(cell)  # Ensure borders are removed for new cells too
                    li_option = getattr(question, f'list_i_option_{chr(96+i)}', '')
                    ii_option = getattr(question, f'list_ii_option_{i}', '')
                    row_cells[0].text = f"{chr(64+i)}. {li_option}" if li_option else ""
                    row_cells[1].text = f"{i}. {ii_option}" if ii_option else ""

                # Clear the original cell content, insert question text and sub-table
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                # p.add_run(clean_text(question.question_text) + "\n").bold = True
                p.add_run(question.question_text + "\n")
                q_row[1]._element.append(sub_table._element)

                # Add 'Codes:' text below the sub-table within the same cell
                p = q_row[1].add_paragraph()
                # p.add_run("\nCodes:").bold = True
                p.add_run("\nCodes:")


            else:
                # Standard question text handling
                q_row[1].text = question_text

            # Merging cells for question text and image
            q_row[1].merge(q_row[2])
            
           # Handling image insertion if available
            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)
                img_io = io.BytesIO()
                pil_img.save(img_io, 'JPEG')
                img_io.seek(0)
                paragraph = q_row[1].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(1.5))


            # ***************************** Table type Questions Start **************************************
            # Check for the presence of table_head_* fields and their corresponding data
            # Initialize sub_table to None
            sub_table = None

            # Data preparation and table creation
            # Data preparation and table creation
            table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
            data_fields = [
                [getattr(question, f'head_a_data{j}', None) for j in range(1, 6)],
                [getattr(question, f'head_b_data{j}', None) for j in range(1, 6)],
                [getattr(question, f'head_c_data{j}', None) for j in range(1, 6)],
                [getattr(question, f'head_d_data{j}', None) for j in range(1, 6)]
            ]

            filtered_heads_data = [
                (head, [data for data in datas if data])
                for head, datas in zip(table_heads, data_fields) 
                if getattr(question, head, None) and any(datas)
            ]

            if filtered_heads_data:
                total_rows = 1 + max(len(datas) for _, datas in filtered_heads_data)
                sub_table = document.add_table(rows=total_rows, cols=len(filtered_heads_data))
                sub_table.style = 'Table Grid'
                hdr_cells = sub_table.rows[0].cells
                for idx, (head, _) in enumerate(filtered_heads_data):
                    hdr_cells[idx].text = getattr(question, head, "")
                    paragraph = hdr_cells[idx].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align header text

                for col_idx, (head, data_list) in enumerate(filtered_heads_data):
                    serial_number = 1
                    for row_idx, data in enumerate(data_list):
                        cell = sub_table.cell(row_idx + 1, col_idx)
                        if head == 'table_head_a':
                            cell.text = f"{serial_number}. {data}"
                            serial_number += 1
                        else:
                            cell.text = data

            if sub_table is not None:
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run(question_text + "\n")
                q_row[1]._element.append(sub_table._element)

            if question.list_i_selection_text:
                p = q_row[1].add_paragraph()
                p.add_run("\n" + question.list_i_selection_text)


            # # Data preparation and table creation
            # table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
            # data_fields = [
            #     [getattr(question, f'head_a_data{j}', None) for j in range(1, 6)],
            #     [getattr(question, f'head_b_data{j}', None) for j in range(1, 6)],
            #     [getattr(question, f'head_c_data{j}', None) for j in range(1, 6)],
            #     [getattr(question, f'head_d_data{j}', None) for j in range(1, 6)]
            # ]

            # filtered_heads_data = [
            #     (head, [data for data in datas if data])
            #     for head, datas in zip(table_heads, data_fields) 
            #     if getattr(question, head, None) and any(datas)
            # ]

            # if filtered_heads_data:
            #     total_rows = 1 + max(len(datas) for _, datas in filtered_heads_data)
            #     sub_table = document.add_table(rows=total_rows, cols=len(filtered_heads_data))
            #     sub_table.style = 'Table Grid'
            #     hdr_cells = sub_table.rows[0].cells
            #     for idx, (head, _) in enumerate(filtered_heads_data):
            #         hdr_cells[idx].text = getattr(question, head, "")

            #     for col_idx, (_, data_list) in enumerate(filtered_heads_data):
            #         for row_idx, data in enumerate(data_list):
            #             cell = sub_table.cell(row_idx + 1, col_idx)
            #             cell.text = data

            # # Check if sub_table was created and append it if so
            # if sub_table is not None:
            #     q_row[1]._element.clear_content()
            #     p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
            #     p.add_run(clean_text(question.question_text) + "\n")
            #     q_row[1]._element.append(sub_table._element)

            # if question.list_i_selection_text:
            #     p = q_row[1].add_paragraph()
            #     p.add_run("\n" + clean_text(question.list_i_selection_text))

            # ***************************** Table type Questions End **************************************

            

            # Additional rows for type, options, solution, and marks
            type_row = table.add_row().cells
            type_row[0].text = 'Type'
            type_row[1].text = question.question_type
            type_row[1].merge(type_row[2])

            # correct_option_text = ""  # Store text of the correct answer
            # for opt in ['a', 'b', 'c', 'd', 'e']:
            #     option_text = getattr(question, f"option_{opt}", None)
            #     if option_text:
            #         opt_row = table.add_row().cells
            #         opt_row[0].text = 'Option'
            #         opt_row[1].text = f"{opt.upper()}. {clean_text(option_text)}"
            #         opt_row[2].text = 'correct' if opt == question.correct_answer_option.lower() else 'incorrect'
            #         if opt == question.correct_answer_option.lower():
            #             correct_option_text = clean_text(option_text)
            correct_option_text = ""  # Store text of the correct answer
            valid_options = ['a', 'b', 'c', 'd', 'e']  # Include 'e' as a valid option, but handle it carefully

            for opt in valid_options:
                option_text = getattr(question, f"option_{opt}", None)  # Get the option text or None if it doesn't exist
                if option_text:  # Check if the option text is not None
                    opt_row = table.add_row().cells
                    opt_row[0].text = 'Option'
                    opt_row[1].text = f"{opt.upper()}. {clean_text(option_text)}"  # Clean and set the option text
                    is_correct = opt == question.correct_answer_option.lower()  # Determine if this option is correct
                    opt_row[2].text = 'correct' if is_correct else 'incorrect'  # Set 'correct' or 'incorrect'
                    if is_correct:
                        correct_option_text = clean_text(option_text)  # Store the correct option text

            # This approach ensures that all options a to e are considered, but only those that exist are processed.

            solution_row = table.add_row().cells
            solution_row[0].text = 'Solution'
            solution_row[1].text = correct_option_text
            solution_row[1].merge(solution_row[2])

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            document.add_paragraph()  # Add space between questions

        document.save(file_path)
        return HttpResponse(f"Document created successfully at: {file_path}")
    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)



def view_generate_document(request):
    return render(request, 'generate_doc.html')